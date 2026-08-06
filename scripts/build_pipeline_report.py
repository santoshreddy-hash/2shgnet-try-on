"""Build a Google Docs-ready technical report for the ear landmark pipeline."""
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "reports"
OUT.mkdir(parents=True, exist_ok=True)
REPORT = OUT / "Ear_Landmark_Virtual_Piercing_Pipeline_Report.docx"
DIAGRAM = OUT / "pipeline_flow.png"

BLACK = "000000"
MUTED = "555555"
GRAY_BORDER = "DADCE0"
LIGHT_GRAY = "F8F9FA"
BLUE = "#1A73E8"
GREEN = "#188038"
RED = "#D93025"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
    ]
    for path in candidates:
        if path and Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def build_diagram(path: Path):
    width, height = 1600, 560
    image = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(image)
    title = font(36, True)
    label = font(24, True)
    detail = font(18)
    d.text((60, 35), "Direct browser frame pipeline", fill="#202124", font=title)
    steps = [
        ("1  Input", "Frozen camera frame", BLUE),
        ("2  YOLO tip", "Ear tip + side", BLUE),
        ("3  Crop box", "Tip-centred ROI", GREEN),
        ("4  Crop", "256 × 256 BGR CHW", GREEN),
        ("5  SHGNet-56", "56 landmarks", RED),
        ("6  One Euro", "Relative-shape filter", "#9334E6"),
        ("7  Final", "Overlay + #56", "#202124"),
    ]
    x, y, box_w, box_h, gap = 50, 170, 200, 230, 20
    for i, (head, sub, color) in enumerate(steps):
        bx = x + i * (box_w + gap)
        d.rounded_rectangle((bx, y, bx + box_w, y + box_h), radius=18, fill="#FFFFFF", outline=color, width=5)
        d.rounded_rectangle((bx + 12, y + 14, bx + box_w - 12, y + 80), radius=10, fill=color)
        d.text((bx + 23, y + 31), head, fill="white", font=label)
        # Simple visual marks distinguish each stage without implying a model result.
        if i == 0:
            d.rectangle((bx + 42, y + 104, bx + 158, y + 182), outline="#5F6368", width=3)
        elif i == 1:
            d.ellipse((bx + 78, y + 112, bx + 122, y + 156), fill=RED)
            d.line((bx + 40, y + 184, bx + 158, y + 184), fill="#5F6368", width=3)
        elif i in (2, 3):
            d.rectangle((bx + 54, y + 108, bx + 146, y + 200), outline=GREEN, width=4)
            if i == 3:
                for p in range(4):
                    d.line((bx + 64, y + 120 + p * 19, bx + 136, y + 120 + p * 19), fill="#9AA0A6", width=2)
        elif i == 4:
            for px, py in [(72, 122), (111, 116), (137, 151), (91, 178), (130, 191)]:
                d.ellipse((bx + px - 5, y + py - 5, bx + px + 5, y + py + 5), fill=RED)
        elif i == 5:
            for px, py in [(72, 122), (111, 116), (137, 151), (91, 178), (130, 191)]:
                d.ellipse((bx + px - 5, y + py - 5, bx + px + 5, y + py + 5), fill="#9334E6")
        else:
            for px, py in [(72, 122), (111, 116), (137, 151), (91, 178), (130, 191)]:
                d.ellipse((bx + px - 5, y + py - 5, bx + px + 5, y + py + 5), fill="#FF6D01")
            d.ellipse((bx + 112, y + 174, bx + 130, y + 192), outline=RED, width=3)
        bbox = d.textbbox((0, 0), sub, font=detail)
        d.text((bx + (box_w - (bbox[2] - bbox[0])) / 2, y + 208), sub, fill="#3C4043", font=detail)
        if i < len(steps) - 1:
            ax = bx + box_w + 2
            d.polygon([(ax, y + 111), (ax + gap - 2, y + 111), (ax + gap - 2, y + 99), (ax + gap + 14, y + 126), (ax + gap - 2, y + 153), (ax + gap - 2, y + 141), (ax, y + 141)], fill="#9AA0A6")
    d.text((60, 465), "Each stage uses the same frozen input frame. If a worker is busy, the frame is skipped instead of queued.", fill="#5F6368", font=detail)
    image.save(path)


def set_font(run, size=11, color=BLACK, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRAY_BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Heading 1", 20, BLACK, 20, 6),
        ("Heading 2", 16, BLACK, 18, 6),
        ("Heading 3", 14, "434343", 16, 4),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = False
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, 26, BLACK, False)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run(subtitle)
        set_font(r, 11, MUTED)


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        set_font(r, 11, BLACK, True)
    r = p.add_run(text)
    set_font(r, 11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        set_font(r, 11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        set_font(r, 11)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, 10, BLACK, True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, 10)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def build_report():
    build_diagram(DIAGRAM)
    doc = Document()
    style_document(doc)
    add_title(
        doc,
        "Ear Landmark & Virtual Piercing Pipeline",
        "Technical implementation report · Direct browser frame pipeline · 5 August 2026",
    )
    add_para(
        doc,
        "This report documents the implemented browser and model pipeline for detecting an ear, predicting 56 ear landmarks, and placing a virtual piercing at landmark #56. The current design follows one direct path per accepted frame: YOLO detects the ear tip, the application builds a tip-centred crop, SHGNet-56 predicts landmarks, One Euro stabilizes landmark shape, and the final overlay is rendered.",
        "Purpose. ",
    )
    h(doc, "Executive summary")
    add_bullets(doc, [
        "The browser pipeline uses YOLO pose for ear-tip detection and SHGNet-56 for 56 landmark heatmaps, including the piercing landmark at index 55 (landmark #56).",
        "Each accepted pipeline run uses one frozen camera image for YOLO, crop creation, and SHGNet inference. This avoids mixing a YOLO result from one frame with a crop from another.",
        "YOLO and SHGNet run in separate web workers. The display loop paints independently; if a worker is busy, a new frame is skipped rather than queued, preventing inference backlog and visible lag.",
        "The same crop geometry, landmark model, landmark index, preprocessing, validation rules, and piercing equation are kept across device classes. Device tuning changes only capture resolution, target FPS, and operating cadence.",
    ])
    h(doc, "Implemented frame pipeline")
    doc.add_picture(str(DIAGRAM), width=Inches(6.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Figure 1. Direct YOLO → crop → SHGNet → One Euro browser pipeline.")
    set_font(r, 9, MUTED, italic=True)

    h(doc, "Step-by-step workflow", 2)
    add_table(doc, ["Step", "Input", "Operation", "Output"], [
        ("1. Input", "Camera image", "Freeze one mirrored camera frame for the complete pipeline run.", "Stable source frame."),
        ("2. YOLO tip", "Frozen frame", "YOLO pose finds the side-profile ear keypoint, person box, nose, and ear confidence.", "Ear tip, side, confidence, face context."),
        ("3. Crop box", "YOLO tip + anatomy", "Estimate pinna height and build a tip-centred square crop using fixed padding.", "Crop centre and side length."),
        ("4. Crop", "Frozen frame + crop box", "Pad outside image bounds, flip left ears when required, resize to 256 × 256, and convert RGB canvas pixels to BGR CHW input.", "SHGNet-ready tensor."),
        ("5. SHGNet", "256 × 256 tensor", "SHGNet-56 produces 56 heatmaps. Soft-argmax converts each heatmap to an image coordinate.", "56 full-frame landmarks."),
        ("6. One Euro", "Tip-relative landmarks", "Filter only the shape relative to the detected tip. The tip itself remains directly driven by the latest YOLO result.", "Stable relative landmark shape."),
        ("7. Final", "YOLO tip + filtered shape", "Reconstruct landmark positions, draw the crop box and landmarks, and highlight landmark #56 as the piercing location.", "Final AR overlay."),
    ], [1000, 1750, 4200, 2410])

    h(doc, "Core technical rules")
    add_table(doc, ["Rule", "Implemented value / behavior"], [
        ("Crop padding", "1.65 × estimated pinna height; crop remains tip-centred with a small medial/downward adjustment."),
        ("Landmark model", "SHGNet-56; input 1 × 3 × 256 × 256; output 56 heatmaps."),
        ("Piercing landmark", "Landmark #56, represented by zero-based array index 55."),
        ("Piercing equation", "piercing = latest accepted YOLO tip + filtered relative landmark[55]."),
        ("Preprocessing", "Canvas RGBA → BGR → YCrCb luminance equalization → normalized BGR CHW."),
        ("Orientation", "Left ear crops are horizontally flipped for SHGNet consistency, then points are unflipped before remapping."),
        ("Display policy", "Show the final box and landmarks only after a valid YOLO ear and valid SHGNet geometry are available."),
    ], [2200, 7160])

    h(doc, "YOLO ear detection", 1)
    add_para(doc, "YOLO pose is responsible for establishing whether a usable ear is present. The browser selects the ear with the stronger keypoint confidence, rejects ambiguous dual-ear/frontal views, and uses nose-to-ear distance when available to reject non-side-profile detections.")
    add_bullets(doc, [
        "Ear keypoint and box confidence thresholds are 0.25 in the current geometry gate.",
        "The crop is only generated after the side-profile validation passes.",
        "The implementation uses the most recent accepted YOLO result for the final landmark anchor. No state machine, optical-flow tracker, or kinematic predictor is in the direct frame pipeline.",
        "YOLO is scheduled whenever the direct pipeline is free; the target is at least one attempt every two display frames, subject to device inference capacity.",
    ])

    h(doc, "Crop-box and preprocessing design", 1)
    add_para(doc, "The crop uses anatomical cues rather than a fixed pixel rectangle. Pinna height is estimated from YOLO face/ear context, then multiplied by the fixed crop padding. The crop centre is slightly shifted toward the face and downward so the lobe and the piercing region remain inside the square.")
    add_bullets(doc, [
        "A gray-padded square crop allows the ROI to touch the edge of the camera image without distorting the ear.",
        "The crop is resized to 256 × 256 only after padding and orientation handling.",
        "The preprocessing path is intentionally aligned with the Python training path to minimize train/inference mismatch.",
        "A first-lock flip retry is permitted only when the initial landmark geometry is clearly invalid, minimizing additional latency.",
    ])

    h(doc, "SHGNet-56 landmark inference", 1)
    add_para(doc, "SHGNet-56 predicts a heatmap for each landmark. The browser applies local soft-argmax around each heatmap peak, returning sub-pixel 256-space coordinates. Coordinates are unflipped when necessary and remapped through the crop transform to the full camera image.")
    add_table(doc, ["Validation gate", "Purpose"], [
        ("Mean heatmap score", "Rejects weak SHGNet output below the configured minimum score (0.07)."),
        ("Landmark span and aspect", "Rejects collapsed, over-large, or implausible landmark clouds."),
        ("Tip-to-cloud relationship", "Ensures the YOLO tip lies near the expected ear landmark cloud."),
        ("Piercing geometry", "Checks that the #56 point remains anatomically plausible in the lower-ear region."),
    ], [3000, 6360])

    h(doc, "One Euro filtering", 1)
    add_para(doc, "One Euro filtering is applied to the landmark shape in coordinates relative to the current YOLO tip. Filtering relative shape reduces jitter while allowing the complete landmark cloud to follow a new YOLO ear position without smoothing the tip itself. The current browser defaults are loaded from one_euro_settings.json; device profiles may tune the filter parameters while retaining identical geometry and landmark math.")
    add_table(doc, ["Parameter", "Role"], [
        ("min_cutoff", "Base smoothing cutoff. Higher values react more quickly; lower values stabilize more strongly."),
        ("beta", "Makes the filter more responsive during motion, reducing visible trailing."),
        ("d_cutoff", "Smooths the estimated derivative used for adaptive filtering."),
        ("max_step_px", "Caps abrupt landmark jumps from a poor inference result."),
    ], [2200, 7160])

    h(doc, "Cross-device and no-lag behavior", 1)
    add_para(doc, "The project’s quality invariants do not change by device: the crop formula, crop padding, SHGNet model, landmark count, #56 piercing index, preprocessing, and acceptance geometry are shared. Device adaptation is confined to computational controls.")
    add_table(doc, ["Control", "High-performance behavior", "Lower-compute behavior", "Invariant"], [
        ("Camera", "Higher rung, typically 640 × 360 or above.", "Camera ladder may reduce to 480 × 270, 352 × 198, or 320 × 180.", "Same crop and model input geometry."),
        ("Display", "Target around 25 FPS.", "Target is reduced only when necessary to stay responsive.", "Display loop remains independent of inference."),
        ("YOLO / SHG", "Runs whenever the worker is free.", "Runs whenever the worker is free; busy frames are skipped.", "No inference queue; latest completed result wins."),
        ("One Euro", "Can use a less aggressive responsiveness preset.", "Can increase responsiveness to prevent visual trailing at lower inference cadence.", "Filters relative shape only."),
    ], [1500, 2450, 2450, 2960])

    h(doc, "Observed validation results", 1)
    add_para(doc, "The annotated ear-pose validation run shows strong in-domain results. The external EarVN set has no 56-landmark ground truth, so its figures are diagnostic proxy measures rather than true NME/PCK accuracy.")
    add_table(doc, ["Evaluation", "Result", "Interpretation"], [
        ("Stage 3 best validation NME", "0.00472", "Very low normalized landmark error on the annotated validation set."),
        ("Stage 3 piercing error", "1.15 px", "Best recorded #56 point error in the training validation output."),
        ("Stage 3 PCK@0.05", "0.9977", "High in-domain landmark hit rate at the 5% threshold."),
        ("EarVN inference success", "500 / 500", "All selected external images completed inference."),
        ("EarVN geometric validity", "96.2%", "Proxy check for reasonable landmark-cloud geometry."),
        ("EarVN confidence ≥ 0.5", "68.6%", "External-domain confidence remains a key improvement opportunity."),
        ("EarVN flip-consistency PCK@0.05", "73.8%", "Self-consistency is useful for diagnosis but is not ground-truth accuracy."),
    ], [2350, 1600, 5410])

    h(doc, "Browser implementation and operation", 1)
    add_numbered(doc, [
        "Start the local browser demo with npm start from the repository root.",
        "Open the local page and click Load models. The app initializes separate workers for YOLO and SHGNet-56 and selects an execution provider (WebGPU when available, otherwise WASM).",
        "Click Start live cam and grant camera permission. Use a clear single-ear side profile for the initial lock.",
        "The HUD should show the direct sequence: YOLO → crop → SHGNet → One Euro. The final overlay appears only after both models return a valid result.",
        "If performance drops, the camera ladder can lower capture resolution. The pipeline does not queue work, so control remains responsive even when inference is slower than the display loop.",
    ])

    h(doc, "Current implementation status", 1)
    add_bullets(doc, [
        "Direct, serialized same-frame pipeline is implemented in web/infer.js as runFramePipeline().",
        "The browser status explicitly identifies the active sequence as YOLO → crop → SHGNet → One Euro.",
        "The browser app was verified to load both model workers successfully in the local demo.",
        "Live camera performance still requires testing across real high-end, mid-range, and low-end devices because browser execution providers and camera timing differ by hardware and browser.",
    ])

    h(doc, "Risks and recommended next steps", 1)
    add_bullets(doc, [
        "Collect a held-out, landmark-labelled cross-device test set. EarVN proxy results cannot replace ground-truth testing.",
        "Measure actual YOLO and SHG completion rate on representative devices to verify the one-per-two-frame YOLO target under real load.",
        "Record per-device inference latency, overlay error, frame drops, and confidence-gate rejections in the HUD or a structured telemetry log.",
        "Review side-profile thresholds using difficult camera cases: dim light, motion blur, glasses/hair occlusion, different skin tones, and partial ears.",
        "Use the same fixed geometry and validation thresholds for all devices; only tune camera rung and filter responsiveness after test evidence supports the change.",
    ])

    h(doc, "Source artifacts", 1)
    add_bullets(doc, [
        "Browser pipeline: web/infer.js and web/ear_geometry.js.",
        "Device settings: performance_profiles.json and one_euro_settings.json.",
        "Training validation: outputs/train_results.json and docs/SHGNet56_Training_Report_allsplits_46x.md.",
        "External diagnostic evaluation: outputs/earvn_test_500/metrics.json.",
    ])

    doc.core_properties.title = "Ear Landmark & Virtual Piercing Pipeline"
    doc.core_properties.subject = "Technical implementation report"
    doc.core_properties.author = "Codex"
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    build_report()
